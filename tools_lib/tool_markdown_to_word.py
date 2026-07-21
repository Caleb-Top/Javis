"""Javis自创: markdown_to_word"""
TOOL_NAME="markdown_to_word"
TOOL_DESC="将Markdown文件转换为Word文档。保留标题、列表、代码块格式。参数: md_path(输入路径), output_path(可选输出路径)"
TOOL_CATEGORY="document_convert"
TOOL_PARAMS={"type":"object","properties":{},"required":[]}

def handler(**kwargs):
    import os, sys, importlib, json
    from pathlib import Path


    def convert(md_path=None, output_path=None, **kw):
        if not md_path or not os.path.exists(md_path):
            return {"success": False, "output": "", "message": f"文件不存在: {md_path}"}
        if not output_path:
            output_path = str(Path(md_path).with_suffix('.docx'))
        try:
            import markdown
            from docx import Document
            from docx.shared import Pt, Inches, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from bs4 import BeautifulSoup

            # 读取Markdown并转换为HTML
            with open(md_path, 'r', encoding='utf-8') as f:
                md_text = f.read()

            html = markdown.markdown(md_text, extensions=['extra', 'codehilite', 'tables'])
            soup = BeautifulSoup(html, 'html.parser')

            doc = Document()

            for element in soup.children:
                if element.name == 'h1':
                    doc.add_heading(element.get_text(), level=1)
                elif element.name == 'h2':
                    doc.add_heading(element.get_text(), level=2)
                elif element.name == 'h3':
                    doc.add_heading(element.get_text(), level=3)
                elif element.name == 'p':
                    doc.add_paragraph(element.get_text())
                elif element.name == 'ul':
                    for li in element.find_all('li'):
                        p = doc.add_paragraph(li.get_text(), style='List Bullet')
                elif element.name == 'ol':
                    for li in element.find_all('li'):
                        p = doc.add_paragraph(li.get_text(), style='List Number')
                elif element.name == 'pre':
                    code = element.get_text()
                    p = doc.add_paragraph()
                    run = p.add_run(code)
                    run.font.name = 'Courier New'
                    run.font.size = Pt(9)
                elif element.name == 'table':
                    rows = element.find_all('tr')
                    if rows:
                        table = doc.add_table(rows=len(rows), cols=len(rows[0].find_all(['td','th'])))
                        for i, row in enumerate(rows):
                            cells = row.find_all(['td','th'])
                            for j, cell in enumerate(cells):
                                table.rows[i].cells[j].text = cell.get_text()

            doc.save(output_path)
            return {"success": True, "output": output_path, "message": f"Markdown已转换为Word: {output_path}"}
        except Exception as e:
            return {"success": False, "output": "", "message": f"转换失败: {str(e)}"}
    

    result = convert(**kwargs)
    return json.dumps(result, ensure_ascii=False)
