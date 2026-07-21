"""Javis自创: word_to_pdf"""
TOOL_NAME="word_to_pdf"
TOOL_DESC="将Word文档转换为PDF。使用python-docx+reportlab引擎。参数: docx_path(输入路径), output_path(可选输出路径)"
TOOL_CATEGORY="document_convert"
TOOL_PARAMS={"type":"object","properties":{},"required":[]}

def handler(**kwargs):
    import os, sys, importlib, json
    from pathlib import Path


    def convert(docx_path=None, output_path=None, **kw):
        if not docx_path or not os.path.exists(docx_path):
            return {"success": False, "output": "", "message": f"文件不存在: {docx_path}"}
        if not output_path:
            output_path = str(Path(docx_path).with_suffix('.pdf'))
        try:
            from docx import Document
            from reportlab.lib.pagesizes import A4
            from reportlab.pdfgen import canvas
            from reportlab.lib.units import mm

            doc = Document(docx_path)
            c = canvas.Canvas(output_path, pagesize=A4)
            width, height = A4
            y = height - 30*mm

            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    continue
                if y < 30*mm:
                    c.showPage()
                    y = height - 30*mm
                c.setFont("Helvetica", 11)
                # 处理换行
                lines = []
                while text:
                    if len(text) > 80:
                        lines.append(text[:80])
                        text = text[80:]
                    else:
                        lines.append(text)
                        text = ""
                for line in lines:
                    c.drawString(30*mm, y, line)
                    y -= 6*mm

            # 处理表格
            for table in doc.tables:
                if y < 50*mm:
                    c.showPage()
                    y = height - 30*mm
                for row in table.rows:
                    row_text = " | ".join(cell.text for cell in row.cells)
                    c.drawString(30*mm, y, row_text[:100])
                    y -= 5*mm
                y -= 5*mm

            c.save()
            return {"success": True, "output": output_path, "message": f"Word已转换为PDF: {output_path}"}
        except Exception as e:
            return {"success": False, "output": "", "message": f"转换失败: {str(e)}"}
    

    # 从kwargs获取参数
    result = convert(**kwargs)
    return json.dumps(result, ensure_ascii=False)
