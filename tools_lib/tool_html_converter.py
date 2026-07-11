"""Javis自创: html_converter — HTML网页转PDF/DOCX/MD"""
TOOL_NAME="html_converter"
TOOL_DESC="HTML网页/文件转换为PDF/DOCX/Markdown。支持本地HTML文件和URL。参数: html_path(HTML文件路径或URL), output_format(pdf/docx/md,默认pdf), output_path(可选输出路径)"
TOOL_CATEGORY="document_convert"
TOOL_PARAMS={"type":"object","properties":{},"required":[]}

def handler(**kwargs):
    import os, json, re
    from pathlib import Path

    def convert(html_path=None, output_format="pdf", output_path=None, **kw):
        if not html_path:
            return {"success": False, "output": "", "message": "请提供HTML文件路径或URL"}

        # 判断是URL还是本地文件
        is_url = html_path.startswith(('http://', 'https://'))
        if not is_url and not os.path.exists(html_path):
            return {"success": False, "output": "", "message": f"文件不存在: {html_path}"}

        if not output_path:
            base = "webpage"
            if is_url:
                base = re.sub(r'[^a-zA-Z0-9]', '_', html_path)[:50]
            else:
                base = Path(html_path).stem
            output_path = str(Path(os.getcwd()) / f"{base}.{output_format.lower()}")

        try:
            from bs4 import BeautifulSoup

            # 获取HTML内容
            if is_url:
                import urllib.request
                req = urllib.request.Request(html_path, headers={"User-Agent": "Mozilla/5.0"})
                with urllib.request.urlopen(req, timeout=30) as resp:
                    html_content = resp.read().decode('utf-8', errors='replace')
            else:
                with open(html_path, 'r', encoding='utf-8', errors='replace') as f:
                    html_content = f.read()

            soup = BeautifulSoup(html_content, 'html.parser')

            # 移除脚本和样式
            for tag in soup(['script', 'style', 'nav', 'footer', 'iframe']):
                tag.decompose()

            # 提取标题
            title_tag = soup.find('title')
            title = title_tag.get_text().strip() if title_tag else "未命名网页"

            # 提取正文
            body = soup.find('body') or soup

            # 提取结构化内容
            structured = []
            for element in body.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p', 'li', 'pre', 'table', 'blockquote', 'img']):
                tag = element.name
                text = element.get_text().strip()

                if tag.startswith('h'):
                    level = int(tag[1])
                    structured.append(('heading', level, text))
                elif tag == 'p':
                    if text:
                        structured.append(('paragraph', 0, text))
                elif tag == 'li':
                    if text:
                        structured.append(('list_item', 0, f"• {text}"))
                elif tag == 'pre':
                    if text:
                        structured.append(('code', 0, text))
                elif tag == 'table':
                    rows = []
                    for tr in element.find_all('tr'):
                        cells = [td.get_text().strip() for td in tr.find_all(['td', 'th'])]
                        if cells:
                            rows.append(cells)
                    if rows:
                        structured.append(('table', 0, rows))
                elif tag == 'blockquote':
                    if text:
                        structured.append(('quote', 0, text))
                elif tag == 'img':
                    alt = element.get('alt', '')
                    src = element.get('src', '')
                    if alt:
                        structured.append(('image', 0, f"[图片: {alt}]"))

            if not structured:
                # 后备方案：纯文本
                text = body.get_text('\n', strip=True)
                structured = [('paragraph', 0, p) for p in text.split('\n') if p.strip()]

            output_format = output_format.lower()
            if output_format == 'md':
                with open(output_path, 'w', encoding='utf-8') as f:
                    f.write(f"# {title}\n\n")
                    for stype, level, content in structured:
                        if stype == 'heading':
                            f.write(f"\n{'#' * level} {content}\n\n")
                        elif stype == 'paragraph':
                            f.write(f"{content}\n\n")
                        elif stype == 'list_item':
                            f.write(f"{content}\n")
                        elif stype == 'code':
                            f.write(f"\n```\n{content}\n```\n\n")
                        elif stype == 'table':
                            if content and len(content) > 0:
                                f.write("| " + " | ".join(str(c) for c in content[0]) + " |\n")
                                f.write("| " + " | ".join("---" for _ in content[0]) + " |\n")
                                for row in content[1:]:
                                    f.write("| " + " | ".join(str(c) for c in row) + " |\n")
                                f.write("\n")
                        elif stype == 'quote':
                            f.write(f"> {content}\n\n")
                        elif stype == 'image':
                            f.write(f"{content}\n\n")

            elif output_format == 'pdf':
                from reportlab.lib.pagesizes import A4
                from reportlab.pdfgen import canvas
                from reportlab.lib.units import mm
                from reportlab.pdfbase import pdfmetrics
                from reportlab.pdfbase.ttfonts import TTFont

                cjk_font = "Helvetica"
                for fp, fn in [("C:/Windows/Fonts/simhei.ttf", "SimHei"),
                               ("C:/Windows/Fonts/msyh.ttc", "MicrosoftYaHei")]:
                    if os.path.exists(fp):
                        try:
                            pdfmetrics.registerFont(TTFont(fn, fp))
                            cjk_font = fn
                            break
                        except:
                            pass

                c = canvas.Canvas(output_path, pagesize=A4)
                w, h = A4
                y = h - 25*mm
                c.setFont(cjk_font, 16)
                c.drawString(20*mm, y, title)
                y -= 15*mm

                for stype, level, content in structured:
                    if y < 20*mm:
                        c.showPage()
                        y = h - 25*mm

                    if stype == 'heading':
                        c.setFont(cjk_font, max(8, 16 - level * 2))
                        c.drawString(20*mm, y, content)
                        y -= (18 - level) * mm
                    elif stype in ('paragraph', 'list_item', 'quote'):
                        c.setFont(cjk_font, 9)
                        text = str(content)
                        while text:
                            if y < 15*mm:
                                c.showPage()
                                c.setFont(cjk_font, 9)
                                y = h - 25*mm
                            c.drawString(20*mm, y, text[:100])
                            text = text[100:]
                            y -= 5*mm
                    elif stype == 'code':
                        c.setFont("Courier", 8)
                        for line in str(content).split('\n'):
                            if y < 15*mm:
                                c.showPage()
                                c.setFont("Courier", 8)
                                y = h - 25*mm
                            c.drawString(20*mm, y, line[:120])
                            y -= 4*mm
                c.save()

            elif output_format == 'docx':
                from docx import Document
                from docx.shared import Pt
                doc = Document()
                doc.add_heading(title, 0)
                for stype, level, content in structured:
                    if stype == 'heading':
                        doc.add_heading(content, level=min(level, 3))
                    elif stype == 'paragraph':
                        doc.add_paragraph(str(content))
                    elif stype == 'list_item':
                        doc.add_paragraph(str(content), style='List Bullet')
                    elif stype == 'code':
                        p = doc.add_paragraph()
                        run = p.add_run(str(content))
                        run.font.name = 'Courier New'
                        run.font.size = Pt(9)
                    elif stype == 'table':
                        if isinstance(content, list) and len(content) > 0:
                            table = doc.add_table(rows=len(content), cols=len(content[0]))
                            table.style = 'Light Grid Accent 1'
                            for i, row in enumerate(content):
                                for j, cell_val in enumerate(row):
                                    table.rows[i].cells[j].text = str(cell_val)
                    elif stype == 'quote':
                        p = doc.add_paragraph(str(content))
                        p.paragraph_format.left_indent = Pt(20)
                        p.runs[0].italic = True if p.runs else False
                doc.save(output_path)

            size_kb = os.path.getsize(output_path) / 1024
            return {"success": True, "output": output_path,
                    "message": f"HTML已转换为{output_format.upper()} ({size_kb:.1f}KB): {output_path}",
                    "title": title}

        except Exception as e:
            return {"success": False, "output": "", "message": f"转换失败: {str(e)}"}

    result = convert(**kwargs)
    return json.dumps(result, ensure_ascii=False)
