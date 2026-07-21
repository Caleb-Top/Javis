"""Javis自创: excel_to_word"""
TOOL_NAME="excel_to_word"
TOOL_DESC="将Excel数据转换为Word文档(含表格)。参数: excel_path(输入路径), output_path(可选输出路径), sheet_name(可选工作表)"
TOOL_CATEGORY="document_convert"
TOOL_PARAMS={"type":"object","properties":{},"required":[]}

def handler(**kwargs):
    import os, sys, importlib, json
    from pathlib import Path


    def convert(excel_path=None, output_path=None, sheet_name=None, **kw):
        if not excel_path or not os.path.exists(excel_path):
            return {"success": False, "output": "", "message": f"文件不存在: {excel_path}"}
        if not output_path:
            output_path = str(Path(excel_path).with_suffix('.docx'))
        try:
            import pandas as pd
            from docx import Document
            from docx.shared import Pt, Inches

            doc = Document()
            doc.add_heading('Excel数据报表', 0)

            if sheet_name:
                dfs = {sheet_name: pd.read_excel(excel_path, sheet_name=sheet_name)}
            else:
                dfs = pd.read_excel(excel_path, sheet_name=None)

            for sn, df in dfs.items():
                doc.add_heading(f'工作表: {sn}', level=1)

                table = doc.add_table(rows=len(df)+1, cols=len(df.columns))
                table.style = 'Light Grid Accent 1'

                for j, col in enumerate(df.columns):
                    table.rows[0].cells[j].text = str(col)

                for i in range(len(df)):
                    for j in range(len(df.columns)):
                        table.rows[i+1].cells[j].text = str(df.iloc[i, j])

            doc.save(output_path)
            return {"success": True, "output": output_path, "message": f"Excel已转换为Word: {output_path}"}
        except Exception as e:
            return {"success": False, "output": "", "message": f"转换失败: {str(e)}"}
    

    # 从kwargs获取参数
    result = convert(**kwargs)
    return json.dumps(result, ensure_ascii=False)
