"""Javis自创: csv_converter — CSV格式转换器"""
TOOL_NAME="csv_converter"
TOOL_DESC="CSV文件格式转换器。支持CSV↔XLSX/PDF/DOCX/JSON/Markdown/TXT。参数: csv_path(输入路径), target_format(目标格式), output_path(可选输出路径), delimiter(分隔符,默认逗号)"
TOOL_CATEGORY="document_convert"
TOOL_PARAMS={"type":"object","properties":{},"required":[]}

def handler(**kwargs):
    import os, json
    from pathlib import Path

    def convert(csv_path=None, target_format="xlsx", output_path=None, delimiter=",", **kw):
        if not csv_path or not os.path.exists(csv_path):
            return {"success": False, "output": "", "message": f"文件不存在: {csv_path}"}
        if not output_path:
            output_path = str(Path(csv_path).with_suffix(f'.{target_format.lower()}'))

        try:
            import pandas as pd
            import numpy as np

            # 尝试多种编码
            df = None
            for enc in ['utf-8', 'utf-8-sig', 'gbk', 'gb2312', 'latin-1']:
                try:
                    df = pd.read_csv(csv_path, delimiter=delimiter, encoding=enc)
                    break
                except:
                    continue
            if df is None:
                return {"success": False, "output": "", "message": "无法读取CSV文件（编码错误）"}

            # 清理列名
            df.columns = [str(c).strip() for c in df.columns]

            target_format = target_format.lower().replace('.', '')
            if target_format == 'xlsx':
                df.to_excel(output_path, index=False, engine='openpyxl')
            elif target_format == 'pdf':
                from reportlab.lib.pagesizes import A4, landscape
                from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph
                from reportlab.lib import colors
                from reportlab.lib.styles import getSampleStyleSheet

                # 截断大表格
                display_df = df.head(100) if len(df) > 100 else df
                doc = SimpleDocTemplate(output_path, pagesize=landscape(A4) if len(df.columns) > 6 else A4)
                data = [display_df.columns.tolist()] + display_df.astype(str).values.tolist()
                t = Table(data, repeatRows=1)
                t.setStyle(TableStyle([
                    ('BACKGROUND', (0,0), (-1,0), colors.HexColor('#4472C4')),
                    ('TEXTCOLOR', (0,0), (-1,0), colors.white),
                    ('ALIGN', (0,0), (-1,-1), 'CENTER'),
                    ('FONTSIZE', (0,0), (-1,-1), 8),
                    ('GRID', (0,0), (-1,-1), 0.5, colors.grey),
                    ('ROWBACKGROUNDS', (0,1), (-1,-1), [colors.white, colors.HexColor('#D9E2F3')]),
                ]))
                doc.build([t])
            elif target_format == 'docx':
                from docx import Document
                from docx.shared import Pt
                doc = Document()
                doc.add_heading(f'CSV数据: {Path(csv_path).stem}', 0)
                doc.add_paragraph(f'{len(df)} 行 × {len(df.columns)} 列')
                # 限制行数
                display_df = df.head(200) if len(df) > 200 else df
                table = doc.add_table(rows=len(display_df)+1, cols=len(display_df.columns))
                table.style = 'Light Grid Accent 1'
                for j, col in enumerate(display_df.columns):
                    table.rows[0].cells[j].text = str(col)
                for i in range(len(display_df)):
                    for j in range(len(display_df.columns)):
                        table.rows[i+1].cells[j].text = str(display_df.iloc[i, j])
                if len(df) > 200:
                    doc.add_paragraph(f'... (仅显示前200行，共{len(df)}行)')
                doc.save(output_path)
            elif target_format == 'json':
                df.to_json(output_path, orient='records', force_ascii=False, indent=2)
            elif target_format == 'md':
                lines = []
                lines.append(f"# {Path(csv_path).stem}\n")
                lines.append(f"*{len(df)} 行 × {len(df.columns)} 列*\n")
                lines.append("| " + " | ".join(str(c) for c in df.columns) + " |")
                lines.append("| " + " | ".join("---" for _ in df.columns) + " |")
                for _, row in df.head(500).iterrows():
                    lines.append("| " + " | ".join(str(v).replace('|','\\|') for v in row.values) + " |")
                if len(df) > 500:
                    lines.append(f"\n*...仅显示前500行*\n")
                with open(output_path, 'w', encoding='utf-8') as f:
                    f.write('\n'.join(lines))
            elif target_format == 'txt':
                df.to_csv(output_path, sep='\t', index=False)
            else:
                return {"success": False, "output": "", "message": f"不支持 CSV→{target_format}"}

            size_kb = os.path.getsize(output_path) / 1024
            return {"success": True, "output": output_path,
                    "message": f"CSV已转换 ({len(df)}行×{len(df.columns)}列, {size_kb:.1f}KB): {output_path}"}
        except Exception as e:
            return {"success": False, "output": "", "message": f"转换失败: {str(e)}"}

    result = convert(**kwargs)
    return json.dumps(result, ensure_ascii=False)
