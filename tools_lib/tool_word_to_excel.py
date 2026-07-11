"""Javis自创: word_to_excel"""
TOOL_NAME="word_to_excel"
TOOL_DESC="从Word文档中提取表格和文本到Excel。参数: docx_path(输入路径), output_path(可选输出路径)"
TOOL_CATEGORY="document_convert"
TOOL_PARAMS={"type":"object","properties":{},"required":[]}

def handler(**kwargs):
    import os, sys, importlib, json
    from pathlib import Path


    def convert(docx_path=None, output_path=None, **kw):
        if not docx_path or not os.path.exists(docx_path):
            return {"success": False, "output": "", "message": f"文件不存在: {docx_path}"}
        if not output_path:
            output_path = str(Path(docx_path).with_suffix('.xlsx'))
        try:
            from docx import Document
            import pandas as pd

            doc = Document(docx_path)

            with pd.ExcelWriter(output_path, engine='openpyxl') as writer:
                # 提取表格
                tables_found = 0
                for i, table in enumerate(doc.tables):
                    data = []
                    for row in table.rows:
                        data.append([cell.text for cell in row.cells])
                    if data:
                        df = pd.DataFrame(data[1:], columns=data[0] if len(data) > 1 else None)
                        df.to_excel(writer, sheet_name=f"Table_{i+1}", index=False)
                        tables_found += 1

                # 提取段落文本
                paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
                if paragraphs:
                    df_text = pd.DataFrame(paragraphs, columns=["Text"])
                    df_text.to_excel(writer, sheet_name="Paragraphs", index=False)

            return {"success": True, "output": output_path, "message": f"成功提取 {tables_found} 个表格"}
        except Exception as e:
            return {"success": False, "output": "", "message": f"转换失败: {str(e)}"}
    

    # 从kwargs获取参数
    result = convert(**kwargs)
    return json.dumps(result, ensure_ascii=False)
