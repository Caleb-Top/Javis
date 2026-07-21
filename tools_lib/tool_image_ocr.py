"""Javis自创: image_ocr — 图片OCR文字识别，支持输出TXT/PDF/DOCX"""
TOOL_NAME="image_ocr"
TOOL_DESC="对图片进行OCR文字识别。支持中文/英文/多语言，可输出TXT/PDF/DOCX/Excel格式。参数: image_path(图片路径或目录), output_format(txt/pdf/docx/xlsx,默认txt), output_path(可选输出路径), lang(语言chi_sim+eng,默认chi_sim+eng)"
TOOL_CATEGORY="document_convert"
TOOL_PARAMS={"type":"object","properties":{},"required":[]}

def handler(**kwargs):
    import os, json
    from pathlib import Path

    def convert(image_path=None, output_format="txt", output_path=None, lang="chi_sim+eng", **kw):
        if not image_path or not os.path.exists(image_path):
            return {"success": False, "output": "", "message": f"文件不存在: {image_path}"}

        try:
            import pytesseract
            from PIL import Image

            # 收集图片
            images_to_ocr = []
            if os.path.isdir(image_path):
                supported = ['.png', '.jpg', '.jpeg', '.bmp', '.tiff', '.webp']
                images_to_ocr = sorted([os.path.join(image_path, f) for f in os.listdir(image_path)
                                       if os.path.splitext(f)[1].lower() in supported])
            else:
                images_to_ocr = [image_path]

            if not images_to_ocr:
                return {"success": False, "output": "", "message": "未找到有效的图片文件"}

            all_text = ""
            page_texts = []

            for img_file in images_to_ocr:
                img = Image.open(img_file)
                # OCR识别
                text = pytesseract.image_to_string(img, lang=lang)
                page_texts.append((img_file, text))
                all_text += f"\n--- {os.path.basename(img_file)} ---\n{text}\n"

            if not output_path:
                base = Path(image_path).stem if not os.path.isdir(image_path) else "ocr_result"
                output_path = str(Path(image_path).parent / f"{base}.{output_format.lower()}")

            output_format = output_format.lower()
            if output_format in ['txt', 'text']:
                with open(output_path, 'w', encoding='utf-8') as f:
                    f.write(all_text)
            elif output_format == 'pdf':
                from reportlab.lib.pagesizes import A4
                from reportlab.pdfgen import canvas
                from reportlab.lib.units import mm
                from reportlab.pdfbase import pdfmetrics
                from reportlab.pdfbase.ttfonts import TTFont

                # 注册中文字体
                for fp, fn in [("C:/Windows/Fonts/simhei.ttf", "SimHei"),
                               ("C:/Windows/Fonts/msyh.ttc", "MicrosoftYaHei")]:
                    if os.path.exists(fp):
                        try:
                            pdfmetrics.registerFont(TTFont(fn, fp))
                            cjk_font = fn
                            break
                        except:
                            pass
                else:
                    cjk_font = "Helvetica"

                c = canvas.Canvas(output_path, pagesize=A4)
                w, h = A4
                y = h - 20*mm
                c.setFont(cjk_font, 10)

                for img_file, text in page_texts:
                    c.drawString(15*mm, y, f"--- {os.path.basename(img_file)} ---")
                    y -= 8*mm
                    for line in text.split('\n'):
                        if y < 15*mm:
                            c.showPage()
                            c.setFont(cjk_font, 10)
                            y = h - 20*mm
                        c.drawString(15*mm, y, line[:120])
                        y -= 5*mm
                    y -= 8*mm
                c.save()
            elif output_format == 'docx':
                from docx import Document
                doc = Document()
                doc.add_heading('OCR识别结果', 0)
                for img_file, text in page_texts:
                    doc.add_heading(os.path.basename(img_file), level=1)
                    for line in text.split('\n'):
                        if line.strip():
                            doc.add_paragraph(line.strip())
                doc.save(output_path)
            elif output_format == 'xlsx':
                import pandas as pd
                rows = []
                for img_file, text in page_texts:
                    for line in text.split('\n'):
                        if line.strip():
                            rows.append({"文件": os.path.basename(img_file), "内容": line.strip()})
                pd.DataFrame(rows).to_excel(output_path, index=False, engine='openpyxl')

            total_chars = len(all_text.strip())
            return {"success": True, "output": output_path,
                    "message": f"OCR识别完成: {len(images_to_ocr)}张图片, {total_chars}字符",
                    "text_preview": all_text[:500]}

        except ImportError:
            return {"success": False, "output": "",
                    "message": "pytesseract未安装。请执行: pip install pytesseract 并安装Tesseract-OCR"}
        except Exception as e:
            return {"success": False, "output": "", "message": f"OCR失败: {str(e)}"}

    result = convert(**kwargs)
    return json.dumps(result, ensure_ascii=False)
