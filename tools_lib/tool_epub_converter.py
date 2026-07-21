"""Javis自创: epub_converter — EPUB电子书格式转换"""
TOOL_NAME="epub_converter"
TOOL_DESC="EPUB电子书格式转换器。支持EPUB→TXT/PDF/DOCX/Markdown。参数: epub_path(输入路径), output_format(txt/pdf/docx/md,默认txt), output_path(可选输出路径)"
TOOL_CATEGORY="document_convert"
TOOL_PARAMS={"type":"object","properties":{},"required":[]}

def handler(**kwargs):
    import os, json
    from pathlib import Path

    def convert(epub_path=None, output_format="txt", output_path=None, **kw):
        if not epub_path or not os.path.exists(epub_path):
            return {"success": False, "output": "", "message": f"文件不存在: {epub_path}"}
        if not output_path:
            output_path = str(Path(epub_path).with_suffix(f'.{output_format.lower()}'))

        try:
            import ebooklib
            from ebooklib import epub
            from bs4 import BeautifulSoup

            book = epub.read_epub(epub_path)

            # 提取元数据
            title = book.get_metadata('DC', 'title')
            creator = book.get_metadata('DC', 'creator')
            title_str = title[0][0] if title else Path(epub_path).stem
            author_str = creator[0][0] if creator else "未知作者"

            # 提取所有文本内容
            chapters = []
            for item in book.get_items():
                if item.get_type() == ebooklib.ITEM_DOCUMENT:
                    soup = BeautifulSoup(item.get_content(), 'html.parser')
                    # 提取标题
                    headings = []
                    for h in soup.find_all(['h1', 'h2', 'h3']):
                        headings.append((h.name, h.get_text().strip()))

                    # 提取正文
                    text_parts = []
                    for p in soup.find_all(['p', 'div', 'li', 'pre', 'blockquote']):
                        t = p.get_text().strip()
                        if t:
                            text_parts.append(t)

                    if text_parts or headings:
                        chapters.append({
                            'title': headings[0][1] if headings else '',
                            'headings': headings,
                            'content': text_parts
                        })

            if not chapters:
                return {"success": False, "output": "", "message": "EPUB中未找到可提取的文本内容"}

            output_format = output_format.lower()
            if output_format in ['txt', 'text']:
                with open(output_path, 'w', encoding='utf-8') as f:
                    f.write(f"《{title_str}》— {author_str}\n\n")
                    for i, ch in enumerate(chapters):
                        if ch['title']:
                            f.write(f"\n{'='*60}\n{ch['title']}\n{'='*60}\n\n")
                        for p in ch['content']:
                            f.write(p + '\n\n')

            elif output_format == 'md':
                with open(output_path, 'w', encoding='utf-8') as f:
                    f.write(f"# 《{title_str}》\n\n**作者**: {author_str}\n\n---\n\n")
                    for i, ch in enumerate(chapters):
                        if ch['title']:
                            f.write(f"\n## {ch['title']}\n\n")
                        for p in ch['content']:
                            f.write(p + '\n\n')

            elif output_format == 'pdf':
                from reportlab.lib.pagesizes import A4
                from reportlab.pdfgen import canvas
                from reportlab.lib.units import mm
                from reportlab.pdfbase import pdfmetrics
                from reportlab.pdfbase.ttfonts import TTFont

                cjk_font = "Helvetica"
                for fp, fn in [("C:/Windows/Fonts/simhei.ttf", "SimHei"),
                               ("C:/Windows/Fonts/msyh.ttc", "MicrosoftYaHei")]:
                    if os.path.exists(fp):
                        try:
                            pdfmetrics.registerFont(TTFont(fn, fp))
                            cjk_font = fn
                            break
                        except:
                            pass

                c = canvas.Canvas(output_path, pagesize=A4)
                w, h = A4
                y = h - 25*mm
                c.setFont(cjk_font, 16)
                c.drawString(20*mm, y, title_str)
                y -= 10*mm
                c.setFont(cjk_font, 10)
                c.drawString(20*mm, y, f"作者: {author_str}")
                y -= 15*mm

                for ch in chapters:
                    c.setFont(cjk_font, 9)
                    for p in ch['content']:
                        if y < 20*mm:
                            c.showPage()
                            c.setFont(cjk_font, 9)
                            y = h - 25*mm
                        # 中文换行
                        while p:
                            seg = p[:100]
                            p = p[100:]
                            c.drawString(20*mm, y, seg)
                            y -= 5*mm
                    y -= 8*mm
                c.save()

            elif output_format == 'docx':
                from docx import Document
                doc = Document()
                doc.add_heading(f"《{title_str}》", 0)
                doc.add_paragraph(f"作者: {author_str}")
                for ch in chapters:
                    if ch['title']:
                        doc.add_heading(ch['title'], level=1)
                    for p in ch['content']:
                        if p.strip():
                            doc.add_paragraph(p)
                doc.save(output_path)

            size_kb = os.path.getsize(output_path) / 1024
            return {"success": True, "output": output_path,
                    "message": f"EPUB已转换 ({len(chapters)}章, {size_kb:.1f}KB): {output_path}",
                    "title": title_str, "author": author_str}

        except ImportError:
            return {"success": False, "output": "", "message": "ebooklib未安装。请执行: pip install ebooklib beautifulsoup4"}
        except Exception as e:
            return {"success": False, "output": "", "message": f"转换失败: {str(e)}"}

    result = convert(**kwargs)
    return json.dumps(result, ensure_ascii=False)
