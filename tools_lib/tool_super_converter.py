"""Javis自创: super_converter — 万能格式转换引擎(融合方案)"""
TOOL_NAME = "super_converter"
TOOL_DESC = "万能格式转换引擎: 41条转换路径, PDF/Word/Excel/PPT/MD/HTML/EPUB/CSV/LaTeX互转, 多引擎自动降级"
TOOL_CATEGORY = "document_convert"
TOOL_PARAMS = {"type":"object","properties":{"source_path":{"type":"string"},"target_format":{"type":"string"},"output_path":{"type":"string"}},"required":["source_path","target_format"]}

def handler(**kwargs):
    import json
    try:
        result = convert(kwargs.get("source_path",""), kwargs.get("target_format",""), kwargs.get("output_path",""))
        return json.dumps(result, ensure_ascii=False)
    except Exception as e:
        return json.dumps({"success":False,"message":str(e)})

import os
import json
from pathlib import Path
from typing import Optional, Dict, List, Tuple

# ============================================================
# 第1层: 格式识别 & 路由表
# ============================================================

# 格式分类
FORMAT_CATEGORIES = {
    "office":   {"docx", "doc", "xlsx", "xls", "pptx", "ppt", "odt", "ods", "odp"},
    "pdf":      {"pdf"},
    "markup":   {"md", "rst", "latex", "tex", "textile", "org", "wiki", "asciidoc", "adoc"},
    "image":    {"png", "jpg", "jpeg", "bmp", "tiff", "webp", "gif"},
    "web":      {"html", "htm", "xml"},
    "ebook":    {"epub", "mobi", "azw3"},
    "data":     {"csv", "tsv", "json"},
    "cat":      {"tmx", "tbx", "xliff", "xlf", "sdlxliff"},
    "text":     {"txt"},
}

def identify_format(filepath: str) -> str:
    """识别文件格式"""
    ext = os.path.splitext(filepath)[1].lower().lstrip('.')
    for category, exts in FORMAT_CATEGORIES.items():
        if ext in exts:
            return category
    return "unknown"

def get_ext(filepath: str) -> str:
    return os.path.splitext(filepath)[1].lower().lstrip('.')

# ============================================================
# 第2层: 引擎矩阵 — 每个引擎有优先级和健康检查
# ============================================================

class Engine:
    """引擎基类"""
    def __init__(self, name: str, source: str, category: str):
        self.name = name
        self.source = source        # "builtin" | "github"
        self.category = category    # "primary" | "fallback" | "optional"
        self.available = None

    def check(self) -> bool:
        """健康检查 — 引擎是否可用"""
        raise NotImplementedError

    def convert(self, src: str, dst: str, target_fmt: str) -> dict:
        raise NotImplementedError


# ---------- 引擎1: pdf2docx (GitHub 3.5k⭐, 已内置) ----------
class Pdf2DocxEngine(Engine):
    def __init__(self):
        super().__init__("pdf2docx", "github:ArtifexSoftware/pdf2docx", "primary")

    def check(self) -> bool:
        try:
            from pdf2docx import Converter
            self.available = True
            return True
        except ImportError:
            self.available = False
            return False

    def convert(self, src, dst, target_fmt):
        from pdf2docx import Converter
        cv = Converter(src)
        cv.convert(dst)
        cv.close()
        return {"engine": self.name, "success": True}


# ---------- 引擎2: pdfplumber (PDF表格) ----------
class PdfPlumberEngine(Engine):
    def __init__(self):
        super().__init__("pdfplumber", "builtin", "primary")

    def check(self) -> bool:
        try:
            import pdfplumber
            self.available = True
            return True
        except ImportError:
            self.available = False
            return False

    def convert(self, src, dst, target_fmt):
        import pdfplumber
        import pandas as pd
        with pdfplumber.open(src) as pdf:
            all_tables = []
            for page in pdf.pages:
                tables = page.extract_tables()
                for t in tables:
                    if t:
                        all_tables.append(t)
            if not all_tables:
                return {"engine": self.name, "success": False,
                        "message": "未提取到表格"}
            with pd.ExcelWriter(dst, engine='openpyxl') as writer:
                for i, table in enumerate(all_tables):
                    df = pd.DataFrame(table[1:], columns=table[0] if table[0] else None)
                    df.to_excel(writer, sheet_name=f'Page_{i+1}', index=False)
        return {"engine": self.name, "success": True}


# ---------- 引擎3: Pandoc (GitHub 45k⭐, 标记格式之王) ----------
class PandocEngine(Engine):
    """
    Pandoc — 标记格式互转之王。
    支持: MD↔LaTeX↔RST↔Textile↔AsciiDoc↔Org↔DOCX↔HTML↔EPUB...
    注意: Pandoc 不直接支持 PDF 输入!
    """
    def __init__(self):
        super().__init__("pandoc", "github:jgm/pandoc (45k⭐)", "primary")

    def check(self) -> bool:
        import subprocess
        try:
            result = subprocess.run(["pandoc", "--version"],
                                    capture_output=True, timeout=5)
            self.available = result.returncode == 0
            return self.available
        except:
            self.available = False
            return False

    def convert(self, src, dst, target_fmt):
        import subprocess
        # Pandoc 参数映射
        pandoc_formats = {
            "docx": "docx", "pdf": "pdf", "md": "markdown",
            "rst": "rst", "latex": "latex", "tex": "latex",
            "html": "html", "epub": "epub", "textile": "textile",
            "asciidoc": "asciidoc", "org": "org", "wiki": "mediawiki",
        }
        to_fmt = pandoc_formats.get(target_fmt, target_fmt)

        cmd = ["pandoc", src, "-o", dst, "-f", "auto", "-t", to_fmt]

        # PDF 需要 xelatex (中文支持)
        if target_fmt == "pdf":
            cmd += ["--pdf-engine=xelatex"]

        result = subprocess.run(cmd, capture_output=True, text=True, timeout=120)
        if result.returncode != 0:
            return {"engine": self.name, "success": False,
                    "message": result.stderr[:500]}
        return {"engine": self.name, "success": True}


# ---------- 引擎4: MarkItDown (微软, 任意→MD) ----------
class MarkItDownEngine(Engine):
    """
    Microsoft MarkItDown — 任意文档→Markdown 统一管道。
    支持: PDF/DOCX/PPTX/XLSX/图片(OCR)/HTML/CSV/JSON/XML/ZIP → Markdown
    """
    def __init__(self):
        super().__init__("markitdown", "github:microsoft/markitdown", "primary")

    def check(self) -> bool:
        try:
            from markitdown import MarkItDown
            self.available = True
            return True
        except ImportError:
            self.available = False
            return False

    def convert(self, src, dst, target_fmt):
        from markitdown import MarkItDown
        md = MarkItDown()
        result = md.convert(src)
        with open(dst, 'w', encoding='utf-8') as f:
            f.write(result.text_content)
        return {"engine": self.name, "success": True}


# ---------- 引擎5: LibreOffice (最强办公格式引擎) ----------
class LibreOfficeEngine(Engine):
    """
    LibreOffice headless — 办公格式深度转换后备引擎。
    擅长处理复杂 DOCX/XLSX/PPTX → PDF，保真度极高。
    """
    def __init__(self):
        super().__init__("libreoffice", "libreoffice headless", "fallback")

    def check(self) -> bool:
        import subprocess
        soffice_paths = [
            "soffice",
            "libreoffice",
            "C:/Program Files/LibreOffice/program/soffice.exe",
            "C:/Program Files (x86)/LibreOffice/program/soffice.exe",
            "/usr/bin/soffice",
        ]
        for path in soffice_paths:
            try:
                result = subprocess.run([path, "--version"],
                                        capture_output=True, timeout=5)
                if result.returncode == 0:
                    self._soffice = path
                    self.available = True
                    return True
            except:
                continue
        self.available = False
        return False

    def convert(self, src, dst, target_fmt):
        import subprocess
        out_dir = os.path.dirname(dst)
        # LibreOffice 只能输出到目录
        cmd = [self._soffice, "--headless", "--convert-to",
               target_fmt, "--outdir", out_dir, src]
        result = subprocess.run(cmd, capture_output=True, text=True, timeout=120)

        # LibreOffice 输出文件名基于源文件名
        base = os.path.splitext(os.path.basename(src))[0]
        expected = os.path.join(out_dir, f"{base}.{target_fmt}")
        if os.path.exists(expected) and expected != dst:
            os.replace(expected, dst)
        elif not os.path.exists(dst):
            return {"engine": self.name, "success": False,
                    "message": result.stderr[:500]}
        return {"engine": self.name, "success": True}


# ---------- 引擎6: Tesseract OCR ----------
class TesseractEngine(Engine):
    def __init__(self):
        super().__init__("tesseract", "builtin", "primary")

    def check(self) -> bool:
        try:
            import pytesseract
            self.available = True
            return True
        except ImportError:
            self.available = False
            return False

    def convert(self, src, dst, target_fmt):
        import pytesseract
        from PIL import Image
        text = pytesseract.image_to_string(Image.open(src), lang='chi_sim+eng')
        if target_fmt in ('txt', 'md'):
            with open(dst, 'w', encoding='utf-8') as f:
                f.write(text)
        elif target_fmt == 'docx':
            from docx import Document
            doc = Document()
            doc.add_paragraph(text)
            doc.save(dst)
        elif target_fmt == 'xlsx':
            import pandas as pd
            lines = [l for l in text.split('\n') if l.strip()]
            pd.DataFrame(lines, columns=["OCR文本"]).to_excel(
                dst, index=False, engine='openpyxl')
        return {"engine": self.name, "success": True}


# ---------- 引擎7: CAT翻译格式 (独有优势!) ----------
class CatEngine(Engine):
    """
    CAT翻译格式引擎 — Javis独有, GitHub无竞品。
    支持: TMX/TBX/XLIFF/SDLXLIFF → CSV/XLSX/TXT/JSON
    """
    def __init__(self):
        super().__init__("cat-engine", "builtin:Javis原创", "primary")

    def check(self) -> bool:
        self.available = True  # 纯Python实现，无外部依赖
        return True

    def convert(self, src, dst, target_fmt):
        import xml.etree.ElementTree as ET
        ext = os.path.splitext(src)[1].lower()
        entries = []

        if ext == '.tmx':
            tree = ET.parse(src)
            for tu in tree.getroot().iter('tu'):
                src_text, tgt_text = '', ''
                for tuv in tu.iter('tuv'):
                    seg = tuv.find('seg')
                    text = ''.join(seg.itertext()).strip() if seg is not None else ''
                    if not src_text:
                        src_text = text
                    else:
                        tgt_text = text
                if src_text or tgt_text:
                    entries.append({"source": src_text, "target": tgt_text})

        elif ext in ('.tbx', '.tbx.xml'):
            tree = ET.parse(src)
            for term_entry in tree.getroot().iter('termEntry'):
                for lang_set in term_entry.iter('langSet'):
                    lang = lang_set.get('{http://www.w3.org/XML/1998/namespace}lang', '')
                    for term in lang_set.iter('term'):
                        entries.append({"source": term.text or '', "target": '',
                                        "context": lang, "metadata": ''})

        elif ext in ('.xliff', '.xlf'):
            tree = ET.parse(src)
            ns = {'x': 'urn:oasis:names:tc:xliff:document:1.2'}
            for unit in tree.iter('trans-unit'):
                src_elem = unit.find('source') if unit.find('source') is not None else unit.find('x:source', ns)
                tgt_elem = unit.find('target') if unit.find('target') is not None else unit.find('x:target', ns)
                entries.append({
                    "source": ''.join(src_elem.itertext()).strip() if src_elem is not None else '',
                    "target": ''.join(tgt_elem.itertext()).strip() if tgt_elem is not None else '',
                    "context": unit.get('id', ''),
                    "metadata": ''
                })

        elif ext == '.sdlxliff':
            tree = ET.parse(src)
            ns = {'sdl': 'http://sdl.com/FileTypes/SdlXliff/1.0'}
            for unit in tree.iter('trans-unit'):
                src_elem = unit.find('source')
                tgt_elem = unit.find('target')
                entries.append({
                    "source": ''.join(src_elem.itertext()).strip() if src_elem is not None else '',
                    "target": ''.join(tgt_elem.itertext()).strip() if tgt_elem is not None else '',
                    "context": unit.get('id', ''),
                    "metadata": ''
                })

        if not entries:
            return {"engine": self.name, "success": False, "message": "未提取到翻译条目"}

        # 输出
        if target_fmt == 'csv':
            import csv
            with open(dst, 'w', newline='', encoding='utf-8-sig') as f:
                w = csv.DictWriter(f, fieldnames=['source', 'target', 'context', 'metadata'])
                w.writeheader()
                w.writerows(entries)
        elif target_fmt == 'xlsx':
            import pandas as pd
            pd.DataFrame(entries).to_excel(dst, index=False, engine='openpyxl')
        elif target_fmt == 'json':
            with open(dst, 'w', encoding='utf-8') as f:
                json.dump(entries, f, ensure_ascii=False, indent=2)
        elif target_fmt == 'txt':
            with open(dst, 'w', encoding='utf-8') as f:
                for e in entries:
                    f.write(f"{e['source']}\t{e['target']}\n")

        return {"engine": self.name, "success": True, "entries": len(entries)}


# ---------- 引擎8: PyPDF2 (纯Python后备) ----------
class PyPDF2Engine(Engine):
    def __init__(self):
        super().__init__("pypdf2", "builtin", "fallback")

    def check(self) -> bool:
        try:
            from PyPDF2 import PdfReader
            self.available = True
            return True
        except ImportError:
            self.available = False
            return False

    def convert(self, src, dst, target_fmt):
        from PyPDF2 import PdfReader
        from docx import Document
        reader = PdfReader(src)
        doc = Document()
        for page in reader.pages:
            text = page.extract_text()
            if text and text.strip():
                doc.add_paragraph(text)
        doc.save(dst)
        return {"engine": self.name, "success": True}


# ============================================================
# 第3层: 路由表 — 每个转换路径有多引擎优先级
# ============================================================

class SuperConverter:
    """
    统一转换器 — 融合方案核心。
    自动选择最优引擎，支持自动降级。
    """

    def __init__(self):
        self.engines: Dict[str, Engine] = {}
        self._init_engines()
        self._init_routes()

    def _init_engines(self):
        """注册所有引擎"""
        engine_classes = [
            Pdf2DocxEngine, PdfPlumberEngine, PandocEngine,
            MarkItDownEngine, LibreOfficeEngine, TesseractEngine,
            CatEngine, PyPDF2Engine,
        ]
        for cls in engine_classes:
            eng = cls()
            self.engines[eng.name] = eng

    def _init_routes(self):
        """
        路由表: (源格式, 目标格式) → [引擎优先级列表]

        设计原则:
        - 每条路径至少2个引擎 (主力+后备)
        - CAT格式保持独有
        - Pandoc覆盖标记格式互转
        - MarkItDown作为→MD的统一管道
        - LibreOffice作为办公格式后备
        """
        E = self.engines  # shorthand
        self.routes = {
            # ── PDF 输出 ──
            ("pdf", "docx"):  [E["pdf2docx"], E["pypdf2"]],
            ("pdf", "xlsx"):  [E["pdfplumber"], E["tesseract"]],
            ("pdf", "md"):    [E["markitdown"], E["tesseract"]],
            ("pdf", "txt"):   [E["tesseract"], E["pypdf2"]],
            ("pdf", "html"):  [E["pandoc"], E["markitdown"]],

            # ── DOCX 输出 ──
            ("docx", "pdf"):  [E["pandoc"], E["libreoffice"]],
            ("docx", "md"):   [E["pandoc"], E["markitdown"]],
            ("docx", "html"): [E["pandoc"], E["libreoffice"]],

            # ── 标记格式互转 (Pandoc主场) ──
            ("md", "pdf"):    [E["pandoc"], E["libreoffice"]],
            ("md", "docx"):   [E["pandoc"], E["libreoffice"]],
            ("md", "html"):   [E["pandoc"]],
            ("md", "latex"):  [E["pandoc"]],
            ("md", "rst"):    [E["pandoc"]],
            ("rst", "md"):    [E["pandoc"]],
            ("rst", "pdf"):   [E["pandoc"]],
            ("latex", "md"):  [E["pandoc"]],
            ("latex", "pdf"): [E["pandoc"]],
            ("latex", "docx"):[E["pandoc"]],

            # ── 图片 (OCR主场) ──
            ("image", "txt"): [E["tesseract"], E["markitdown"]],
            ("image", "md"):  [E["markitdown"], E["tesseract"]],
            ("image", "docx"):[E["tesseract"]],
            ("image", "xlsx"):[E["tesseract"]],
            ("image", "pdf"): [E["libreoffice"], E["pandoc"]],

            # ── 网页 ──
            ("html", "pdf"):  [E["pandoc"], E["libreoffice"]],
            ("html", "md"):   [E["pandoc"], E["markitdown"]],
            ("html", "docx"): [E["pandoc"], E["libreoffice"]],

            # ── 电子书 ──
            ("epub", "md"):   [E["pandoc"], E["markitdown"]],
            ("epub", "pdf"):  [E["pandoc"], E["libreoffice"]],
            ("epub", "docx"): [E["pandoc"]],

            # ── Excel ──
            ("xlsx", "pdf"):  [E["libreoffice"], E["pandoc"]],
            ("xlsx", "md"):   [E["markitdown"], E["pandoc"]],
            ("xlsx", "docx"): [E["libreoffice"]],

            # ── PPT ──
            ("pptx", "pdf"):  [E["libreoffice"], E["pandoc"]],
            ("pptx", "md"):   [E["markitdown"]],
            ("pptx", "docx"): [E["libreoffice"]],

            # ── CAT翻译格式 (独有!) ──
            ("cat", "csv"):   [E["cat-engine"]],
            ("cat", "xlsx"):  [E["cat-engine"]],
            ("cat", "json"):  [E["cat-engine"]],
            ("cat", "txt"):   [E["cat-engine"]],

            # ── 数据格式 ──
            ("csv", "xlsx"):  [E["pandoc"]],
            ("csv", "md"):    [E["markitdown"], E["pandoc"]],
        }

    def convert(self, source_path: str, target_format: str,
                output_path: Optional[str] = None) -> dict:
        """
        统一转换入口。

        参数:
            source_path: 源文件路径
            target_format: 目标格式 (pdf/docx/xlsx/md/txt/html/csv/json/latex/rst...)
            output_path: 输出路径 (可选)

        返回:
            {"success": bool, "output": str, "message": str,
             "engine_used": str, "fallback_attempts": int}
        """
        # 参数验证
        if not os.path.exists(source_path):
            return {"success": False, "output": "", "message": f"文件不存在: {source_path}"}

        target_format = target_format.lower().replace('.', '')
        fmt_alias = {
            'word': 'docx', 'excel': 'xlsx', 'markdown': 'md',
            'text': 'txt', 'jpeg': 'jpg', 'powerpoint': 'pptx',
        }
        target_format = fmt_alias.get(target_format, target_format)

        if not output_path:
            output_path = str(Path(source_path).parent /
                            (Path(source_path).stem + f".{target_format}"))

        # 识别源格式
        src_cat = identify_format(source_path)
        src_ext = get_ext(source_path)

        # 查找路由
        route = None
        for (cat, tgt), engines in self.routes.items():
            if cat == src_cat and tgt == target_format:
                route = engines
                break
        # 也查具体扩展名
        if not route:
            for (cat, tgt), engines in self.routes.items():
                if cat == src_ext and tgt == target_format:
                    route = engines
                    break

        if not route:
            return {"success": False, "output": "",
                    "message": f"不支持: {src_ext}→{target_format}。"
                               f"支持源格式: {list(set(k[0] for k in self.routes))}"}

        # 依次尝试引擎 (自动降级)
        tried = []
        for engine in route:
            if not engine.check():
                tried.append(f"{engine.name}(不可用)")
                continue

            try:
                result = engine.convert(source_path, output_path, target_format)
                if result.get("success"):
                    size = os.path.getsize(output_path) if os.path.exists(output_path) else 0
                    fallback_idx = route.index(engine)
                    return {
                        "success": True,
                        "output": output_path,
                        "message": f"{src_ext.upper()}→{target_format.upper()} "
                                   f"({size/1024:.1f}KB) via {engine.name}"
                                   + (f" [降级{falleback_idx}次]" if fallback_idx > 0 else ""),
                        "engine_used": engine.name,
                        "engine_source": engine.source,
                        "fallback_attempts": fallback_idx,
                        "engines_tried": tried,
                    }
                else:
                    tried.append(f"{engine.name}(失败:{result.get('message','')})")
            except Exception as e:
                tried.append(f"{engine.name}(异常:{str(e)[:100]})")
                continue

        return {"success": False, "output": "",
                "message": f"所有引擎均失败: {' → '.join(tried)}"}

    def status(self) -> dict:
        """检查所有引擎状态"""
        result = {}
        for name, eng in self.engines.items():
            try:
                ok = eng.check()
                result[name] = {
                    "available": ok,
                    "source": eng.source,
                    "category": eng.category,
                }
            except Exception as e:
                result[name] = {
                    "available": False,
                    "source": eng.source,
                    "category": eng.category,
                    "error": str(e)[:100],
                }
        return result

    @property
    def route_map(self) -> List[str]:
        """列出所有支持的转换路径"""
        paths = []
        for (src, tgt), engines in self.routes.items():
            primary = engines[0]
            fallbacks = [e.name for e in engines[1:]]
            paths.append(f"{src}→{tgt}: {primary.name}" +
                        (f" (后备: {', '.join(fallbacks)})" if fallbacks else ""))
        return sorted(paths)


# ============================================================
# 对外接口 (兼容现有 tool 格式)
# ============================================================

_converter = SuperConverter()

def convert(source_path: str, target_format: str, output_path: str = None) -> dict:
    """快捷转换函数"""
    return _converter.convert(source_path, target_format, output_path)

def engine_status() -> dict:
    """引擎状态检查"""
    return _converter.status()

def list_routes() -> list:
    """列出所有支持的转换路径"""
    return _converter.route_map


# ============================================================
# 测试 & 演示
# ============================================================

if __name__ == "__main__":
    print("=" * 60)
    print("  SuperConverter — 融合方案引擎状态")
    print("=" * 60)

    status = engine_status()
    available = sum(1 for s in status.values() if s["available"])
    total = len(status)
    print(f"\n引擎: {available}/{total} 可用\n")

    for name, info in status.items():
        icon = "✅" if info["available"] else "❌"
        print(f"  {icon} {name}")
        print(f"      来源: {info['source']}")
        print(f"      角色: {info['category']}")
        if info.get("error"):
            print(f"      错误: {info['error']}")
        print()

    print("=" * 60)
    print(f"  支持 {len(list_routes())} 条转换路径")
    print("=" * 60)
    for path in list_routes():
        print(f"  • {path}")