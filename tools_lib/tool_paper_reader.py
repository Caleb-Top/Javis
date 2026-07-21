"""Javis自创: paper_reader v2 — 学术论文深度阅读与结构化提取"""
TOOL_NAME="paper_reader"
TOOL_DESC="学术论文阅读与格式转换工具。支持PDF论文，提取标题/作者/摘要/参考文献等结构化信息，可转换为Word/Excel/Markdown。参数: pdf_path(论文路径), output_format(word/excel/md,默认md), output_path(输出路径)"
TOOL_CATEGORY="document_convert"
TOOL_PARAMS={"type":"object","properties":{},"required":[]}

def handler(**kwargs):
    import os, json, re
    from pathlib import Path

    def convert(pdf_path=None, output_format="md", output_path=None, **kw):
        if not pdf_path or not os.path.exists(pdf_path):
            return {"success": False, "output": "", "message": f"文件不存在: {pdf_path}"}

        try:
            import pdfplumber
            import re

            # 提取文本（优先使用fitz获得更好结构）
            full_text = ""
            page_texts = []
            try:
                import fitz
                doc = fitz.open(pdf_path)
                for page in doc:
                    text = page.get_text()
                    page_texts.append(text)
                    full_text += text + "\n"
                doc.close()
            except:
                with pdfplumber.open(pdf_path) as pdf:
                    for page in pdf.pages:
                        text = page.extract_text() or ""
                        page_texts.append(text)
                        full_text += text + "\n"

            lines = [l.strip() for l in full_text.split('\n') if l.strip()]

            # === 智能结构化提取 ===
            title = lines[0] if lines else "Unknown Title"

            # 作者提取（多种模式）
            authors = ""
            author_patterns = [
                r'([A-Z][a-z]+ [A-Z]\. ?[A-Z]?[a-z]*)',  # John D. Smith
                r'([A-Z][a-z]+ [A-Z][a-z]+)',              # John Smith
                r'([A-Z]\. ?[A-Z]\. [A-Z][a-z]+)',          # J.D. Smith
            ]
            for i, line in enumerate(lines[:30]):
                for pat in author_patterns:
                    found = re.findall(pat, line)
                    if len(found) >= 2 and not any(kw in line.lower() for kw in ['university','institute','department','college','school','abstract','introduction']):
                        if len(line) < 200:
                            authors = line
                            break
                if authors:
                    break
            if not authors:
                # 尝试从元数据获取
                try:
                    import fitz
                    doc = fitz.open(pdf_path)
                    meta = doc.metadata
                    if meta.get('author'):
                        authors = meta['author']
                    doc.close()
                except:
                    pass

            # 摘要提取（多种模式）
            abstract = ""
            abstract_start = -1
            for i, line in enumerate(lines):
                if re.search(r'^\s*abstract\s*$|^abstract[:\-—]', line, re.IGNORECASE):
                    abstract_start = i + 1
                    break
            if abstract_start == -1:
                # 查找Abstract关键词在行内
                for i, line in enumerate(lines):
                    if 'abstract' in line.lower() and len(line) < 120:
                        abstract_start = i
                        break

            if abstract_start >= 0:
                for i in range(abstract_start, min(abstract_start + 30, len(lines))):
                    if re.search(r'^\s*(introduction|1\.?\s|keywords?|index terms|key words)', lines[i], re.IGNORECASE):
                        break
                    if len(lines[i]) > 20 and not lines[i].startswith('©') and not lines[i].startswith('Fig'):
                        abstract += lines[i] + " "

            # 关键词提取
            keywords = ""
            for i, line in enumerate(lines):
                if re.search(r'(keywords?|index terms|key words)[:\-—]?\s', line, re.IGNORECASE):
                    kw_line = re.sub(r'(keywords?|index terms|key words)[:\-—]?\s*', '', line, flags=re.IGNORECASE)
                    keywords = kw_line.strip()
                    break

            # 参考文献提取
            references = []
            ref_start = -1
            for i, line in enumerate(lines):
                if re.search(r'^\s*references?\s*$|^\s*bibliography\s*$|^REFERENCES\s*$', line):
                    ref_start = i + 1
                    break
            if ref_start < 0:
                # 从后往前搜
                for i in range(len(lines)-1, max(len(lines)-50, 0), -1):
                    if re.search(r'references?|bibliography', lines[i], re.IGNORECASE):
                        ref_start = i + 1
                        break

            if ref_start >= 0:
                for i in range(ref_start, min(ref_start + 80, len(lines))):
                    line = lines[i].strip()
                    if line and len(line) > 10:
                        # 匹配编号引用
                        if re.match(r'^\[\d+\]|^\d+\.?\s', line):
                            references.append(line)
                        elif references and len(line) > 20:
                            # 续行
                            references[-1] += " " + line

            # === 章节结构识别 ===
            sections = []
            section_pattern = re.compile(
                r'^(\d+\.?\s*(?:INTRODUCTION|RELATED|METHOD|EXPERIMENT|RESULT|DISCUSSION|CONCLUSION|'
                r'ABSTRACT|REFERENCES?|BACKGROUND|APPENDIX|ACKNOWLEDGMENT|EVALUATION|ANALYSIS|'
                r'引言|绪论|相关|方法|实验|结果|讨论|结论|参考文献|附录|致谢|背景|分析|评估).*)',
                re.IGNORECASE
            )
            for line in lines:
                m = section_pattern.match(line)
                if m:
                    sections.append(m.group(1))

            if not sections:
                # 备用：纯数字章节
                for line in lines:
                    if re.match(r'^(?:I+V*|[1-9]\d*)\.?\s+[A-Z][a-z]+', line):
                        if len(line) < 120:
                            sections.append(line)

            # === 输出 ===
            if not output_path:
                output_path = str(Path(pdf_path).with_suffix(f'.{output_format.lower()}'))

            output_format = output_format.lower()
            if output_format == "md":
                with open(output_path, 'w', encoding='utf-8') as f:
                    f.write(f"# {title}\n\n")
                    if authors:
                        f.write(f"**作者**: {authors}\n\n")
                    if keywords:
                        f.write(f"**关键词**: {keywords}\n\n")
                    f.write(f"## 摘要\n\n{abstract.strip()}\n\n" if abstract.strip() else "## 摘要\n\n未识别到摘要\n\n")

                    if sections:
                        f.write(f"## 章节结构\n\n")
                        for s in sections[:20]:
                            f.write(f"- {s}\n")
                        f.write("\n")

                    if references:
                        f.write(f"## 参考文献 ({len(references)}条)\n\n")
                        for i, ref in enumerate(references[:50]):
                            f.write(f"{i+1}. {ref}\n")
                        f.write("\n")

                    # 全文附录
                    f.write(f"## 全文\n\n")
                    f.write(full_text[:10000])
                    if len(full_text) > 10000:
                        f.write(f"\n\n*...（全文共{len(full_text)}字符，仅显示前10000字符）*")

            elif output_format == "word":
                from docx import Document
                from docx.shared import Pt, Inches

                doc = Document()
                doc.add_heading(title, 0)
                if authors:
                    doc.add_paragraph(f"作者: {authors}")
                if keywords:
                    doc.add_paragraph(f"关键词: {keywords}")

                doc.add_heading('摘要', level=1)
                doc.add_paragraph(abstract.strip() if abstract.strip() else "未识别到摘要")

                if sections:
                    doc.add_heading('章节结构', level=1)
                    for s in sections[:20]:
                        doc.add_paragraph(s, style='List Bullet')

                if references:
                    doc.add_heading(f'参考文献 ({len(references)}条)', level=1)
                    for i, ref in enumerate(references[:50]):
                        doc.add_paragraph(f"{i+1}. {ref}")

                doc.add_heading('全文内容', level=1)
                doc.add_paragraph(full_text[:5000])
                if len(full_text) > 5000:
                    doc.add_paragraph(f"...（全文共{len(full_text)}字符）")

                doc.save(output_path)

            elif output_format == "excel":
                import pandas as pd
                with pd.ExcelWriter(output_path, engine='openpyxl') as writer:
                    # 基本信息
                    info_data = {
                        "项目": ["标题", "作者", "关键词", "摘要", "总字符数", "参考文献数"],
                        "内容": [title, authors.strip(), keywords, abstract.strip(), len(full_text), len(references)]
                    }
                    pd.DataFrame(info_data).to_excel(writer, sheet_name="基本信息", index=False)

                    # 章节
                    if sections:
                        pd.DataFrame({"章节": sections[:50]}).to_excel(writer, sheet_name="章节结构", index=False)

                    # 参考文献
                    if references:
                        pd.DataFrame({
                            "序号": range(1, len(references[:50])+1),
                            "参考文献": references[:50]
                        }).to_excel(writer, sheet_name="参考文献", index=False)

                    # 全文
                    text_lines = [l for l in full_text.split('\n') if l.strip()]
                    pd.DataFrame({
                        "行号": range(1, len(text_lines[:500])+1),
                        "内容": text_lines[:500]
                    }).to_excel(writer, sheet_name="全文(前500行)", index=False)

            return {
                "success": True, "output": output_path,
                "message": f"论文分析完成: '{title[:60]}'",
                "title": title,
                "authors": authors.strip()[:200],
                "keywords": keywords,
                "abstract_preview": abstract.strip()[:300],
                "section_count": len(sections),
                "reference_count": len(references),
                "total_chars": len(full_text)
            }

        except Exception as e:
            import traceback
            return {"success": False, "output": "", "message": f"论文处理失败: {str(e)}"}

    result = convert(**kwargs)
    return json.dumps(result, ensure_ascii=False)
